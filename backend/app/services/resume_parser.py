import io
import logging
from typing import Optional
from fastapi import HTTPException

try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        fitz = None

try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger("skillforge.resume_parser")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF using PyMuPDF.

    Raises HTTPException if PDF is image-only/scanned without extractable text.
    """
    if not fitz:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF library is not installed on the server."
        )

    try:
        text_parts = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts).strip()
        
        if not full_text:
            raise HTTPException(
                status_code=400,
                detail="Unable to extract text from PDF. The document appears to be scanned or image-only, which requires OCR processing."
            )
            
        return full_text
    except HTTPException:
        raise
    except Exception as e:
        logger.error("PDF parsing error: %s", str(e))
        raise HTTPException(
            status_code=400,
            detail=f"Corrupted or invalid PDF file: {str(e)}"
        )


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX paragraphs and tables."""
    if not docx:
        raise HTTPException(
            status_code=500,
            detail="python-docx library is not installed on the server."
        )

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts).strip()
        if not full_text:
            raise HTTPException(
                status_code=400,
                detail="Unable to extract text from DOCX file. The file may be empty or contain unsupported elements."
            )
        return full_text
    except HTTPException:
        raise
    except Exception as e:
        logger.error("DOCX parsing error: %s", str(e))
        raise HTTPException(
            status_code=400,
            detail=f"Corrupted or invalid DOCX file: {str(e)}"
        )


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from raw text file supporting utf-8 and fallback encodings."""
    try:
        try:
            return file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1").strip()
    except Exception as e:
        logger.error("TXT decoding error: %s", str(e))
        raise HTTPException(
            status_code=400,
            detail="Failed to decode text file. Ensure it is valid text."
        )


def parse_file_to_text(file_name: str, file_bytes: bytes) -> str:
    """Master text extractor validating file format and returning clean extracted text."""
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    lower_name = file_name.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx") or lower_name.endswith(".doc"):
        if lower_name.endswith(".doc"):
            raise HTTPException(
                status_code=400,
                detail="Legacy .doc format is not supported. Please convert your resume to .docx, .pdf, or .txt format."
            )
        return extract_text_from_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload a PDF (.pdf), Word (.docx), or Text (.txt) file."
        )
